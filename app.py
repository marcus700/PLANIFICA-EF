import io
import re
import time
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
import google.generativeai as genai
import streamlit as st

# Importamos la base de datos oficial CNEB completa y la función de ciclo
from cneb_datos import CNEB_PRIMARIA, obtener_ciclo_primaria

# --- LISTAS DE OPCIONES ---
GRADOS = [
    "1° de Primaria",
    "2° de Primaria",
    "3° de Primaria",
    "4° de Primaria",
    "5° de Primaria",
    "6° de Primaria"
]

DURACIONES_UNIDAD = [
    "4 Semanas (4 sesiones)",
    "5 Semanas (5 sesiones)",
    "6 Semanas (6 sesiones)",
    "8 Semanas (8 sesiones)"
]

COMPETENCIAS_EF = [
    "Se desenvuelve de manera autónoma a través de su motricidad",
    "Asume una vida saludable",
    "Interactúa a través de sus habilidades sociomotrices"
]

# 1. Configuración de la Plataforma Streamlit
st.set_page_config(page_title="Generador CNEB Primaria", page_icon="🏃‍♂️", layout="centered")

st.title("🏃 Generador CNEB - Educación Física (Primaria)")
st.subheader("Con Estándares y Desempeños oficiales del MINEDU (1° a 6° de Primaria)")
st.write("Herramienta inteligente para diseñar tus documentos curriculares al instante utilizando la base de datos CNEB oficial.")

# Configuramos la clave API
try:
    api_key = st.secrets["GEMINI_API_KEY"]
    genai.configure(api_key=api_key)
except Exception:
    st.error("🔑 No se encontró la variable GEMINI_API_KEY en los Secrets de Streamlit.")

# FUNCIÓN DE GENERACIÓN CON GEMINI OPTIMIZADA (Alta precisión y sin recortes)
def generar_con_gemini(prompt, instrucciones_sistema=""):
    modelos_a_probar = [
        "gemini-1.5-pro",        # Intentar primero con el modelo Pro (más detallado)
        "models/gemini-1.5-pro",
        "gemini-1.5-flash",
        "models/gemini-1.5-flash"
    ]
    
    # Configuración para evitar resúmenes y permitir documentos largos completos
    config_generacion = {
        "temperature": 0.2,           # Menor aleatoriedad para máxima fidelidad al CNEB
        "max_output_tokens": 8192     # Límite amplio para generar todas las tablas de las 10 secciones
    }
    
    ultimo_error = None
    
    for nombre_modelo in modelos_a_probar:
        try:
            model = genai.GenerativeModel(model_name=nombre_modelo, system_instruction=instrucciones_sistema)
            return model.generate_content(prompt, generation_config=config_generacion)
        except Exception as e:
            error_msg = str(e)
            if "429" in error_msg or "quota" in error_msg.lower():
                time.sleep(5)
                try:
                    return model.generate_content(prompt, generation_config=config_generacion)
                except Exception as ex:
                    ultimo_error = ex
                    continue
            else:
                ultimo_error = e
                continue
                
    raise ultimo_error

# FUNCIONES AUXILIARES PARA FORMATEO DE WORD (.docx)
def add_formatted_text(paragraph, text):
    """Procesa negritas **texto** dentro de un párrafo de python-docx"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def render_markdown_table(doc, lines):
    """Convierte líneas de tabla Markdown a tabla nativa de MS Word con bordes y formato"""
    rows = []
    for line in lines:
        if '---' in line:
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)

    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_value in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                add_formatted_text(p, cell_value)
                if i == 0:  # Encabezado de la tabla en negrita
                    for run in p.runs:
                        run.bold = True

def crear_archivo_word(texto_contenido):
    """Genera un archivo Word (.docx) procesando títulos, listas, negritas y tablas Markdown"""
    doc = Document()
    
    # Configuración de márgenes estándar de 1 pulgada
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = texto_contenido.split('\n')
    in_table = False
    table_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            render_markdown_table(doc, table_lines)
            in_table = False
            table_lines = []

        if not stripped:
            continue

        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(27, 94, 32)
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(46, 125, 50)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(27, 94, 32)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)

    if in_table:
        render_markdown_table(doc, table_lines)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Creación de las 3 Pestañas de Navegación
tab1, tab2, tab3 = st.tabs([
    "📂 Crear Unidad de Aprendizaje", 
    "📝 Crear Sesión de Aprendizaje", 
    "📊 Crear Rúbrica de Evaluación"
])

# ==============================================================================
# PESTAÑA 1: UNIDADES DE APRENDIZAJE
# ==============================================================================
with tab1:
    st.write("Estructura una unidad didáctica completa alineada al CNEB con su matriz de planificación y secuencia de sesiones.")
    with st.form("form_unidad"):
        col1, col2 = st.columns(2)
        with col1:
            num_unidad_u = st.text_input("🔢 N° de Unidad:", value="Unidad N° 04", key="u_num")
            grado_u = st.selectbox("🏫 Grado de Primaria:", GRADOS, key="u1")
            duracion_u = st.selectbox("⏱️ Duración de la Unidad:", DURACIONES_UNIDAD, key="u2")
            ie_u = st.text_input("🏢 Nombre de la I.E.:", value='N.° 22314 "Vicenta Aquije de Huamán"', key="u_ie")
            directora_u = st.text_input("👤 Directora:", value="Prof. Luisa Ruth Aronés Herrera", key="u_dir")
        with col2:
            docente_u = st.text_input("👨‍🏫 Docente de Ed. Física:", value="Mario A. García Torres", key="u_doc")
            fechas_u = st.text_input("📅 Fechas / Periodo:", value="22 de junio al 17 de julio de 2026", key="u_fechas")
            producto_u = st.text_input("🏆 Producto de la Unidad:", value="Festival Lúdico-Motor Peruanito", key="u_prod")

        problema_u = st.text_area(
            "📋 Describe la problemática o tema central a abordar:",
            placeholder="Ej. Dificultades de coordinación motriz, orientación espacial, juegos tradicionales y hábitos de higiene",
            value="Dificultades de coordinación motriz, orientación espacial en el patio, juegos tradicionales y falta de hábitos de higiene personal al concluir la actividad física.",
            key="u3"
        )
        boton_unidad = st.form_submit_button("📂 Generar Unidad de Aprendizaje en Word")

    if boton_unidad and problema_u:
        with st.spinner("Escribiendo la unidad didáctica oficial CNEB..."):
            try:
                ciclo_u = obtener_ciclo_primaria(grado_u)
                
                # Extraemos la base de datos oficial del CNEB para el ciclo y grado seleccionados
                cneb_contexto = ""
                for comp_nombre, comp_data in CNEB_PRIMARIA.items():
                    est_texto = comp_data["estandares"].get(ciclo_u, "")
                    des_lista = comp_data["desempenos"].get(grado_u, [])
                    cneb_contexto += f"\n\nCOMPETENCIA: {comp_nombre}\nESTÁNDAR OFICIAL ({ciclo_u}):\n{est_texto}\nDESEMPEÑOS OFICIALES ({grado_u}):\n" + "\n".join(des_lista)

                # ==============================================================================
                # PROMPT MAESTRO COMPLETO INTEGRADO
                # ==============================================================================
                prompt_maestro = f"""
Actúa como un especialista en currículo educativo peruano y docente experto en el área de Educación Física para Educación Básica Regular (CNEB). 

Tu tarea es elaborar una UNIDAD DE APRENDIZAJE completa, rigurosa y alineada al Currículo Nacional (CNEB), siguiendo estrictamente la estructura y reglas del modelo proporcionado a continuación.

DATOS OFICIALES EXTRAÍDOS DEL CNEB PARA UTILIZAR EN ESTA UNIDAD ({grado_u} - {ciclo_u}):
{cneb_contexto}

DATOS PARA LA GENERACIÓN:
- N° de Unidad: {num_unidad_u}
- Ciclo / Grados: {ciclo_u} - {grado_u}
- Nombre de la IE: {ie_u}
- Nombre del Docente: {docente_u}
- Nombre del Director(a): {directora_u}
- Duración / Fechas: {duracion_u} - {fechas_u}
- Tema central / Problemática a abordar: {problema_u}
- Producto de la Unidad: {producto_u}

---

ESTRUCTURA OBLIGATORIA DE LA UNIDAD DE APRENDIZAJE:

1. TÍTULO DE LA UNIDAD
- Debe ser motivador, entre comillas y redactado en función al desarrollo de competencias motrices o lúdicas.

2. II. DATOS INFORMATIVOS
- IE, Directora, Profesor, Ciclo, Grado y Sección, Duración.

3. III. SITUACIÓN SIGNIFICATIVA
- Contextualizar la realidad de los estudiantes relacionada con la problemática.
- Incluir un dato cuantitativo/cualitativo del problema (ej. "solo el X% logra...").
- Plantear 3 preguntas retadoras/desafiantes asociadas a la solución.
- Proponer la estrategia pedagógica para resolver el reto (juegos, circuitos, festivales, etc.).

4. IV. PRODUCTO DE LA UNIDAD
- Describir un desempeño práctico o un producto tangible/demostrable claro donde el estudiante aplique lo aprendido.

5. V. ENFOQUES TRANSVERSALES
- Seleccionar 2 enfoques transversales del CNEB.
- Especificar: Enfoque Transversal, Valor(es) y Acciones o Actitudes Observables adaptadas a Educación Física.

6. VI. COMPETENCIAS TRANSVERSALES
- Incluir "Gestiona su aprendizaje de manera autónoma" y "Se desenvuelve en entornos virtuales generados por las TIC" con sus respectivas Capacidades y Desempeños aplicados al área.

7. VII. ESTÁNDARES, COMPETENCIAS Y CAPACIDADES DEL ÁREA DE EDUCACIÓN FÍSICA
- Incluir las 3 competencias del área con sus capacidades y estándares completos del ciclo correspondiente ({ciclo_u}):
  * Competencia 1: Se desenvuelve de manera autónoma a través de su motricidad.
  * Competencia 2: Asume una vida saludable.
  * Competencia 3: Interactúa a través de sus habilidades sociomotrices.

8. VIII. MATRIZ DE PLANIFICACIÓN (Formato Tabla detallado por cada sesión)
Estructura de la tabla por cada sesión/actividad:
- En la parte superior de cada bloque/actividad, incluir la fila del ESTÁNDAR COMPLETO del CNEB correspondiente a la competencia evaluada, redactado de manera íntegra (sin modificar ni alterar su texto original), pero RESALTANDO EN NEGRITA la parte específica que se trabaja/evalúa en esa actividad.
- Columnas de la Matriz:
  1. Sesión N.° y Título de la sesión
  2. Competencia / Capacidad
  3. Desempeño (Redactado de manera COMPLETA tal cual aparece en el CNEB del ciclo, pero RESALTANDO EN NEGRITA tanto la parte del desempeño utilizada como las palabras/términos agregados para su precisión y contextualización).
  4. Criterios de evaluación (mínimo 3 por sesión, precisados y observables)
  5. Evidencia y Producto
  6. Instrumento de evaluación (Lista de cotejo, Escala de valoración, Guía de observación, etc.)

*NOTA: NO incluir la columna "Propósito" en la Matriz de Planificación.*

9. IX. SECUENCIA DE SESIONES (Formato Tabla)
Para cada una de las sesiones planificadas, detalla:
- N° y Título de la actividad
- Propósito de la actividad (explicando la secuencia metodológica: calentamiento, desarrollo motriz/deportivo, hábitos de higiene y reflexión).
- Representación gráfica (descripción breve del gráfico o material visual sugerido).

10. X. RECURSOS
- Recursos para el Docente (Normativa CNEB, documentos vigentes, materiales).
- Recursos para el Estudiante (Kit de aseo, ropa deportiva, materiales específicos).
- Fecha y espacio para firmas (Directora y Docente de Educación Física).

Asegúrate de cumplir estrictamente la regla de negritas para el Estándar y Desempeño en la Matriz de Planificación.

GENERA AHORA LA UNIDAD DE APRENDIZAJE COMPLETA Y DETALLADA SIGUIENDO EXACTAMENTE LAS 10 SECCIONES.
"""

                response = generar_con_gemini(prompt_maestro)
                
                resultado_u = response.text
                st.success("¡Unidad Curricular CNEB generada con éxito!")
                st.markdown(resultado_u)
                archivo_word_u = crear_archivo_word(resultado_u)
                st.download_button(
                    label="📄 Descargar Unidad de Aprendizaje en Word (.docx)", 
                    data=archivo_word_u, 
                    file_name=f"Unidad_Aprendizaje_EF_{grado_u.replace(' ', '_')}.docx", 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"❌ Error al generar la unidad: {str(e)}")

# ==============================================================================
# PESTAÑA 2: SESIONES DE APRENDIZAJE
# ==============================================================================
with tab2:
    st.write("Completa los datos oficiales requeridos por el MINEDU para tu clase diaria.")
    with st.form("form_sesion"):
        grado_s = st.selectbox("🏫 Grado de Primaria:", GRADOS, key="s1")
        duracion_s = st.number_input("⏱️ Duración (minutos):", value=90, min_value=30, max_value=180, key="s2")
        competencia_s = st.selectbox("🎯 Competencia CNEB:", COMPETENCIAS_EF, key="s3")
        tema_s = st.text_input("⚽ Tema / Propósito Motriz:", value="Coordinación dinámica general y equilibrio en desplazamientos", key="s4")
        materiales_s = st.text_input("🧱 Materiales Disponibles:", value="Conos, aros, pelotas, tizas, kit de aseo", key="s5")
        boton_sesion = st.form_submit_button("🚀 Generar Sesión de Aprendizaje CNEB Primaria")

    if boton_sesion and tema_s and materiales_s:
        with st.spinner("Escribiendo la sesión oficial CNEB..."):
            try:
                ciclo_s = obtener_ciclo_primaria(grado_s)
                estandar_real = CNEB_PRIMARIA.get(competencia_s, {}).get("estandares", {}).get(ciclo_s, "Estándar general.")
                desempenos_reales = CNEB_PRIMARIA.get(competencia_s, {}).get("desempenos", {}).get(grado_s, ["Desempeños generales."])
                desempenos_texto = "\n".join(desempenos_reales)
                
                instrucciones = f"""Actúa como un Asistente Pedagógico experto en Educación Física para Primaria bajo el enfoque oficial del CNEB del MINEDU de Perú.
Estás redactando para {grado_s} ({ciclo_s}).

DATOS OFICIALES CNEB:
- Estándar: "{estandar_real}"
- Desempeños disponibles:
{desempenos_texto}

Estructura la Sesión incluyendo de forma ordenada:
1. Datos Informativos (IE, Docente, Grado, Fecha, Área: Educación Física, Duración: {duracion_s} min).
2. Propósitos y Evidencias de Aprendizaje (Tabla con: Competencia y Capacidades, Estándar CNEB con **negrita** en la parte aplicada, Desempeño precisado completo con **negrita**, Criterios de Evaluación, Evidencia y Producto, Instrumento: Lista de Cotejo).
3. Enfoques Transversales.
4. Preparación de la Sesión.
5. Secuencia Didáctica (Inicio {round(duracion_s * 0.2)} min, Desarrollo {round(duracion_s * 0.6)} min con progresiones y variaciones, Cierre {round(duracion_s * 0.2)} min con vuelta a la calma, metacognición e higiene personal) redactada en PRIMERA PERSONA y TIEMPO PRESENTE.
6. Anexo: Lista de Cotejo."""

                pedido = f"Diseña una sesión de {duracion_s} minutos para {grado_s} ({ciclo_s}). Tema: {tema_s}. Materiales: {materiales_s}."
                
                response = generar_con_gemini(pedido, instrucciones_sistema=instrucciones)
                
                resultado_s = response.text
                st.success("¡📋 Sesión Generada (MINEDU) con éxito!")
                st.markdown(resultado_s)
                archivo_word = crear_archivo_word(resultado_s)
                st.download_button(
                    label="📄 Descargar Sesión en Word (.docx)", 
                    data=archivo_word, 
                    file_name=f"Sesion_EF_{grado_s.replace(' ', '_')}.docx", 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"❌ Error al generar la sesión: {str(e)}")

# ==============================================================================
# PESTAÑA 3: RÚBRICAS DE EVALUACIÓN
# ==============================================================================
with tab3:
    st.write("Diseña instrumentos de evaluación con criterios claros y descriptores por niveles de logro CNEB.")
    with st.form("form_rubrica"):
        grado_r = st.selectbox("🏫 Grado de Primaria:", GRADOS, key="r1")
        competencia_r = st.selectbox("🎯 Competencia a Evaluar:", COMPETENCIAS_EF, key="r2")
        criterio_r = st.text_input("📊 Desempeño o criterio específico a evaluar:", value="Orientación espacial y equilibrio postural en actividades lúdicas")
        boton_rubrica = st.form_submit_button("📊 Generar Rúbrica en Word")

    if boton_rubrica and criterio_r:
        with st.spinner("Escribiendo la rúbrica oficial CNEB..."):
            try:
                ciclo_r = obtener_ciclo_primaria(grado_r)
                instrucciones_r = f"""Actúa como un Evaluador Pedagógico experto en Educación Física para Primaria bajo los lineamientos del CNEB del MINEDU. 
Diseña una rúbrica analítica en formato tabla estructurada para evaluar a estudiantes de {grado_r} ({ciclo_r}). 
Competencia: {competencia_r}. 
Actividad/Criterio a evaluar: {criterio_r}.
Incluye los cuatro niveles de logro oficiales del CNEB: En Inicio, En Proceso, Logrado y Logro Destacado."""

                pedido_r = f"Crea una rúbrica analítica de evaluación para {grado_r} ({ciclo_r}). Criterio: {criterio_r}"
                
                response = generar_con_gemini(pedido_r, instrucciones_sistema=instrucciones_r)
                
                resultado_r = response.text
                st.success("¡Rúbrica CNEB generada con éxito!")
                st.markdown(resultado_r)
                archivo_word_r = crear_archivo_word(resultado_r)
                st.download_button(
                    label="📄 Descargar Rúbrica en Word (.docx)", 
                    data=archivo_word_r, 
                    file_name=f"Rubrica_EF_{grado_r.replace(' ', '_')}.docx", 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"❌ Error al generar la rúbrica: {str(e)}")
